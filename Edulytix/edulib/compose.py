from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
import os, pickle, sys

def run():
    d = os.path.dirname(__file__)
    data_file = os.path.join(d, "report_data.dat")
    plot_file = os.path.join(d, "performance_plot.png")

    # Exit gracefully if required files do not exist
    if not os.path.exists(data_file) or not os.path.exists(plot_file):
        sys.exit(0)

    from . import gemini_summary as gem

    with open(data_file, "rb") as f:
        data = pickle.load(f)
    doc_path = os.path.join(data[7], f"{data[0]}_Report.docx")

    Image.open(plot_file).rotate(270, expand=True).save(os.path.join(d, "rotated_image.png"))

    def add_section(doc, heading, content, multiline=False):
        p = doc.add_paragraph()
        run = p.add_run(heading)
        run.bold = True
        run.font.size = Pt(13)
        if multiline:
            doc.add_paragraph(content).runs[0].font.size = Pt(12)
        else:
            p.add_run(" " + content).font.size = Pt(12)

    doc = Document()
    for s in doc.sections:
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Inches(0.5)
    
    # Add document title with bold and larger font size
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("EduLytix - Scholastic Report")
    title_run.bold = True
    title_run.font.size = Pt(22)  # Set font size to 22
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for i, h in enumerate(["Student Name:", "Class:", "Section:", "Attendance:"]):
        add_section(doc, h, data[i])
    add_section(doc, "Assignment Completion:", "\n".join(f"{k.title()}: {v}%" for k, v in data[4].items()), multiline=True)
    add_section(doc, "Teacher's Remarks:", data[5], multiline=True)
    
    # AI Summary heading with same style as Teacher's Remarks
    ai_summary_para = doc.add_paragraph()
    ai_summary_run = ai_summary_para.add_run("AI Summary:")
    ai_summary_run.bold = True
    ai_summary_run.font.size = Pt(13)
    doc.add_paragraph(gem.pass_output().strip()).runs[0].font.size = Pt(12)
    
    doc.add_page_break()
    doc.add_paragraph("Subject Wise Performance Graph").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().add_run().add_picture(os.path.join(d, "rotated_image.png"), width=Inches(6))
    doc.save(doc_path)
